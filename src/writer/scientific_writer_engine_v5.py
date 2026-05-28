import pandas as pd
from pathlib import Path
from docx import Document


class ScientificWriterEngineV5:

    def __init__(self, output_dir="outputs"):
        self.output_dir = Path(output_dir)
        self.writer_dir = self.output_dir / "scientific_writer"
        self.writer_dir.mkdir(parents=True, exist_ok=True)

    def find_outputs(self):
        return {
            "review": list(self.output_dir.rglob("athena_scientific_review_v4.docx")),
            "stats": list(self.output_dir.rglob("assumptions_posthoc_results.xlsx")),
            "descriptive": list(self.output_dir.rglob("descriptive_table.xlsx")),
            "similarity": list(self.output_dir.rglob("similarity_summary.xlsx")),
            "umap": list(self.output_dir.rglob("umap_coordinates.xlsx")),
            "hdbscan": list(self.output_dir.rglob("hdbscan_clusters.xlsx"))
        }

    def write_results(self, files):
        paragraphs = []

        paragraphs.append(
            "Os resultados foram analisados de forma automatizada pelo ATHENA, integrando estatística descritiva, testes inferenciais, análises multivariadas e procedimentos de agrupamento, conforme a estrutura do conjunto de dados."
        )

        if files["descriptive"]:
            paragraphs.append(
                "A estatística descritiva permitiu caracterizar a tendência central, a dispersão e a variabilidade das variáveis quantitativas avaliadas."
            )

        if files["stats"]:
            try:
                df = pd.read_excel(files["stats"][0])

                if "global_p" in df.columns:
                    sig = df[df["global_p"].astype(float) < 0.05]

                    if not sig.empty:
                        variables = ", ".join(sig["variable"].astype(str).tolist())
                        paragraphs.append(
                            f"Os testes globais identificaram diferenças estatisticamente significativas nas variáveis: {variables}. Esses achados sugerem diferenças relevantes entre os grupos avaliados."
                        )
                    else:
                        paragraphs.append(
                            "Os testes globais não identificaram diferenças estatisticamente significativas entre os grupos avaliados."
                        )

            except Exception:
                paragraphs.append(
                    "Os testes inferenciais foram executados, porém a leitura automática detalhada dos resultados não pôde ser concluída."
                )

        if files["similarity"]:
            paragraphs.append(
                "As análises de similaridade contribuíram para avaliar a estrutura multivariada dos dados, permitindo comparar amostras ou grupos com base em padrões conjuntos de variáveis."
            )

        if files["umap"]:
            paragraphs.append(
                "A redução dimensional por UMAP foi utilizada para explorar padrões não lineares e possíveis estruturas latentes no conjunto de dados."
            )

        if files["hdbscan"]:
            paragraphs.append(
                "O HDBSCAN foi aplicado para identificar agrupamentos naturais e possíveis observações discrepantes no espaço multivariado."
            )

        return paragraphs

    def write_discussion(self, files):
        paragraphs = []

        paragraphs.append(
            "A integração entre estatística inferencial e análises multivariadas amplia a capacidade interpretativa dos dados, pois permite não apenas testar diferenças entre grupos, mas também reconhecer padrões globais de organização das observações."
        )

        paragraphs.append(
            "Quando diferenças estatísticas são observadas, a interpretação deve considerar o contexto científico do domínio analisado, o tamanho amostral, a estrutura dos grupos e a relevância prática dos efeitos encontrados."
        )

        if files["similarity"]:
            paragraphs.append(
                "Nos conjuntos de dados ambientais, ecológicos, oceanográficos ou citométricos, as análises de similaridade são particularmente úteis porque capturam relações entre amostras considerando múltiplas variáveis simultaneamente."
            )

        if files["umap"] or files["hdbscan"]:
            paragraphs.append(
                "A presença de agrupamentos em métodos como UMAP e HDBSCAN pode indicar perfis latentes, padrões não lineares ou subestruturas que não seriam facilmente identificadas por análises univariadas."
            )

        paragraphs.append(
            "Embora o ATHENA gere interpretações automatizadas, os resultados devem ser avaliados criticamente pelo pesquisador, considerando o desenho do estudo, a qualidade dos dados e as hipóteses científicas originais."
        )

        return paragraphs

    def write_conclusion(self):
        return [
            "O ATHENA v5 permitiu gerar uma síntese científica automatizada dos resultados, integrando análise estatística, exploração multivariada e interpretação metodológica.",
            "Os achados produzidos podem apoiar relatórios institucionais, manuscritos científicos, apresentações acadêmicas e processos de tomada de decisão baseados em dados.",
            "Recomenda-se que a versão final do texto seja revisada pelo pesquisador responsável antes de submissão, publicação ou uso institucional."
        ]

    def write_figure_legends(self, files):
        legends = []

        if files["descriptive"]:
            legends.append(
                "Figura/Tabela 1. Estatística descritiva das variáveis quantitativas avaliadas, incluindo medidas de tendência central, dispersão e coeficiente de variação."
            )

        if files["stats"]:
            legends.append(
                "Figura/Tabela 2. Resultados dos testes de pressupostos, testes globais e pós-testes aplicados às variáveis quantitativas segundo os grupos definidos."
            )

        if files["similarity"]:
            legends.append(
                "Figura/Tabela 3. Análises de similaridade baseadas em matriz multivariada, incluindo ordenação por NMDS e testes PERMANOVA/ANOSIM quando aplicáveis."
            )

        if files["umap"]:
            legends.append(
                "Figura 4. Projeção UMAP das observações, utilizada para explorar padrões não lineares e estruturas latentes no conjunto de dados."
            )

        if files["hdbscan"]:
            legends.append(
                "Figura 5. Agrupamentos identificados pelo algoritmo HDBSCAN, com indicação de clusters e possíveis observações classificadas como ruído."
            )

        return legends

    def generate_document(self, sections):
        doc = Document()

        doc.add_heading("ATHENA Scientific Writer v5", level=1)

        doc.add_paragraph(
            "Texto científico preliminar gerado automaticamente pelo ATHENA v5 a partir dos resultados analíticos disponíveis no projeto."
        )

        doc.add_heading("Resultados", level=2)
        for paragraph in sections["results"]:
            doc.add_paragraph(paragraph)

        doc.add_heading("Discussão", level=2)
        for paragraph in sections["discussion"]:
            doc.add_paragraph(paragraph)

        doc.add_heading("Conclusões", level=2)
        for paragraph in sections["conclusion"]:
            doc.add_paragraph(paragraph)

        doc.add_heading("Legendas sugeridas", level=2)
        for legend in sections["legends"]:
            doc.add_paragraph(legend)

        doc.add_heading("Nota metodológica", level=2)
        doc.add_paragraph(
            "Este texto é uma versão inicial automatizada. A curadoria científica, a checagem dos valores estatísticos e a adequação ao periódico ou relatório final devem ser realizadas pelo pesquisador responsável."
        )

        output = self.writer_dir / "athena_scientific_writer_v5.docx"
        doc.save(output)

        return output

    def run(self):
        print("\n[ATHENA v5] Iniciando Scientific Writer...")

        files = self.find_outputs()

        sections = {
            "results": self.write_results(files),
            "discussion": self.write_discussion(files),
            "conclusion": self.write_conclusion(),
            "legends": self.write_figure_legends(files)
        }

        output = self.generate_document(sections)

        print("[ATHENA v5] Scientific Writer finalizado.")
        print(f"[ATHENA v5] Documento: {output}")

        return output
