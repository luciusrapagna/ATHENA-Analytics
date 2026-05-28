import pandas as pd
from pathlib import Path
from docx import Document


class ScientificReviewerEngineV4:

    def __init__(self, output_dir="outputs"):
        self.output_dir = Path(output_dir)
        self.review_dir = self.output_dir / "scientific_review"
        self.review_dir.mkdir(parents=True, exist_ok=True)

    def find_files(self):
        return {
            "stats": list(self.output_dir.rglob("assumptions_posthoc_results.xlsx")),
            "similarity": list(self.output_dir.rglob("similarity_summary.xlsx")),
            "umap": list(self.output_dir.rglob("umap_coordinates.xlsx")),
            "hdbscan": list(self.output_dir.rglob("hdbscan_clusters.xlsx")),
            "descriptive": list(self.output_dir.rglob("descriptive_table.xlsx"))
        }

    def interpret_stats(self, files):
        texts = []

        for file in files:
            try:
                df = pd.read_excel(file)

                if df.empty:
                    continue

                significant = df[
                    df["global_p"].astype(float) < 0.05
                ] if "global_p" in df.columns else pd.DataFrame()

                texts.append(
                    "A análise dos pressupostos estatísticos e dos testes globais foi executada automaticamente. "
                    "O ATHENA avaliou normalidade, homogeneidade e selecionou o teste global mais adequado para cada variável."
                )

                if not significant.empty:
                    vars_sig = ", ".join(significant["variable"].astype(str).tolist())

                    texts.append(
                        f"Foram identificadas diferenças estatisticamente significativas nas seguintes variáveis: {vars_sig}. "
                        "Esses resultados indicam que pelo menos um dos grupos avaliados apresentou comportamento distinto."
                    )

                else:
                    texts.append(
                        "Não foram identificadas diferenças estatisticamente significativas nos testes globais avaliados."
                    )

            except Exception as e:
                texts.append(f"Não foi possível interpretar estatística: {e}")

        return texts

    def interpret_similarity(self, files):
        texts = []

        for file in files:
            try:
                df = pd.read_excel(file)

                texts.append(
                    "As análises de similaridade foram executadas com base em matriz de distância multivariada, incluindo NMDS, PERMANOVA e ANOSIM quando havia agrupamento disponível."
                )

                cols = " ".join(df.columns.astype(str).tolist()).lower()

                if "permanova" in cols:
                    texts.append(
                        "A PERMANOVA foi utilizada para testar diferenças multivariadas entre grupos, sendo adequada para dados ecológicos, ambientais, oceanográficos e citométricos."
                    )

                if "anosim" in cols:
                    texts.append(
                        "A ANOSIM foi utilizada como análise complementar para avaliar separação entre grupos com base na matriz de similaridade."
                    )

            except Exception as e:
                texts.append(f"Não foi possível interpretar similaridade: {e}")

        return texts

    def interpret_machine_learning(self, umap_files, hdbscan_files):
        texts = []

        if umap_files:
            texts.append(
                "A análise UMAP foi executada para reduzir a dimensionalidade dos dados e revelar padrões não lineares potencialmente não capturados por métodos lineares."
            )

        if hdbscan_files:
            try:
                df = pd.read_excel(hdbscan_files[0])

                if "cluster_hdbscan" in df.columns:
                    clusters = sorted(df["cluster_hdbscan"].dropna().unique().tolist())
                    n_clusters = len([c for c in clusters if c != -1])
                    noise = int((df["cluster_hdbscan"] == -1).sum())

                    texts.append(
                        f"O HDBSCAN identificou {n_clusters} agrupamento(s) principais, com {noise} observação(ões) classificadas como ruído/outlier. "
                        "Esse resultado sugere a presença de padrões latentes no conjunto de dados."
                    )
                else:
                    texts.append(
                        "O HDBSCAN foi executado, mas a coluna de clusters não foi localizada no arquivo de saída."
                    )

            except Exception as e:
                texts.append(f"Não foi possível interpretar HDBSCAN: {e}")

        return texts

    def interpret_descriptive(self, files):
        texts = []

        for file in files:
            try:
                df = pd.read_excel(file)

                texts.append(
                    "A estatística descritiva foi utilizada para caracterizar tendência central, dispersão e variabilidade das variáveis numéricas."
                )

                if "cv_percent" in df.columns:
                    high_cv = df[df["cv_percent"] > 30]

                    if not high_cv.empty:
                        texts.append(
                            "Foram observadas variáveis com coeficiente de variação superior a 30%, indicando heterogeneidade relevante nos dados."
                        )

            except Exception as e:
                texts.append(f"Não foi possível interpretar estatística descritiva: {e}")

        return texts

    def generate_word_review(self, interpretation_sections):
        doc = Document()

        doc.add_heading("ATHENA Scientific Reviewer v4", level=1)

        doc.add_paragraph(
            "Este relatório foi gerado automaticamente pelo ATHENA v4 com base nos resultados estatísticos, multivariados e de similaridade disponíveis na pasta do projeto."
        )

        for title, texts in interpretation_sections.items():
            doc.add_heading(title, level=2)

            if texts:
                for text in texts:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph(
                    "Nenhum resultado específico foi encontrado para esta seção."
                )

        doc.add_heading("Síntese interpretativa", level=2)

        doc.add_paragraph(
            "Em conjunto, os resultados permitem uma leitura integrada do conjunto de dados, combinando estatística inferencial, exploração multivariada, agrupamento e interpretação científica automatizada. "
            "A interpretação deve ser revisada pelo pesquisador responsável antes de uso em relatório institucional, artigo científico ou apresentação acadêmica."
        )

        output = self.review_dir / "athena_scientific_review_v4.docx"

        doc.save(output)

        return output

    def run(self):
        print("\n[ATHENA v4] Iniciando Scientific Reviewer...")

        files = self.find_files()

        interpretation_sections = {
            "Estatística descritiva": self.interpret_descriptive(files["descriptive"]),
            "Pressupostos, testes globais e pós-testes": self.interpret_stats(files["stats"]),
            "Análises de similaridade": self.interpret_similarity(files["similarity"]),
            "Aprendizado de máquina e padrões latentes": self.interpret_machine_learning(
                files["umap"],
                files["hdbscan"]
            )
        }

        output = self.generate_word_review(interpretation_sections)

        print("[ATHENA v4] Scientific Reviewer finalizado.")
        print(f"[ATHENA v4] Relatório: {output}")

        return output
