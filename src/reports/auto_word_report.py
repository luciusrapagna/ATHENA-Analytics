from docx import Document
from pathlib import Path


class AutoWordReport:

    def __init__(self, output_dir="outputs"):

        self.output_dir = Path(output_dir)

        self.reports_dir = (
            self.output_dir / "reports"
        )

        self.reports_dir.mkdir(
            parents=True,
            exist_ok=True
        )

    def generate(
        self,
        interpretation
    ):

        doc = Document()

        doc.add_heading(
            "ATHENA Analytics Report",
            level=1
        )

        doc.add_paragraph(interpretation)

        output = (
            self.reports_dir /
            "athena_report.docx"
        )

        doc.save(output)

        return output
